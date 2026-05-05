"""
Generate performance_analysis.docx for issue #174.
Table columns: Проблема | Что приводит к проблеме | Возможное решение | Потенциальное улучшение
Data sourced from speed_opt/performance_report.md (issue #170 analysis).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROWS = [
    {
        "problem": "Запись «pending» в таблицу cookie_consents при каждом визите",
        "cause": "CookieConsentManager создаёт или читает строку со статусом pending для каждого pageview. Таблица выросла до 2 168 684 строк — самая большая в БД.",
        "solution": "Хранить состояние pending только в браузерном cookie; писать в БД только реальные решения (accepted/rejected). Запустить Artisan-команду для периодического удаления старых pending-строк.",
        "improvement": "TTFB публичных страниц −10–25 %; DB writes −50–70 %; снижение lock waits на горячем индексе таблицы.",
    },
    {
        "problem": "Производительность поиска: LIKE '%слово%' вместо полнотекстового индекса",
        "cause": "FastSearch.php делает LIKE '%...%', который не использует индекс — всегда full scan по 22 000 товаров. В БД уже есть FULLTEXT-индекс items_name_synonyms_fulltext по полям name и synonyms.",
        "solution": "Заменить LIKE-запросы на MATCH ... AGAINST с режимом boolean, добавить точный поиск по vendor_code, barcode, id_1c через индексные запросы.",
        "improvement": "Ускорение endpoint поиска в 3–10 раз; снижение rows examined с ~22 000 до ~десятков на запрос.",
    },
    {
        "problem": "Страница каталога загружает полную карточку товара вместо облегчённого списка",
        "cause": "ItemRepository::catalogue() собирает для каждого товара цены, остатки, скидки, аналоги, атрибуты, характеристики, гайды, изображения и сегменты — всё это нужно только для страницы товара.",
        "solution": "Разделить метод на catalogueList() (только id, uuid, name, vendor_code, минимальная цена, одно изображение) и catalogueDetail() с полным eager-load.",
        "improvement": "TTFB каталога ускоряется в 2–5 раз; значительное снижение потребления памяти PHP-воркера; меньше нагрузки на DB CPU.",
    },
    {
        "problem": "Каталог: get() + groupBy() на PHP вместо SQL пагинации",
        "cause": "Queries/Catalogue.php загружает весь набор товаров категории в память PHP, а группировку выполняет в коде. При крупной категории это тысячи Eloquent-объектов за один HTTP-запрос.",
        "solution": "Перейти на SQL-пагинацию с LIMIT/OFFSET или keyset по items.id; groupBy-логику перенести в SQL GROUP BY или отдельный запрос.",
        "improvement": "Пиковое потребление памяти снижается в разы; p95/p99 latency каталога сокращается; возможность отдавать большие категории без риска out-of-memory.",
    },
    {
        "problem": "Dev-настройки в production: APP_DEBUG=true, file cache/session, sync очереди",
        "cause": ".env.example содержит APP_DEBUG=true, CACHE_DRIVER=file, SESSION_DRIVER=file, QUEUE_CONNECTION=sync, LOG_LEVEL=debug. Эти значения попадают в production и создают постоянный overhead.",
        "solution": "Переключить APP_DEBUG=false, LOG_LEVEL=warning, перейти на Redis для cache/session/queue, запустить php artisan config:cache route:cache view:cache.",
        "improvement": "TTFB −20–40 %; CPU −15–30 %; устранение дискового I/O на каждый запрос; корректная асинхронная обработка экспортов и писем.",
    },
    {
        "problem": "dd() в production-коде: аварийные остановки запросов",
        "cause": "В HomeController::emailFeedback(), app/Models/Model.php, Pipes_/Price/ToCurrency.php и Livewire/Cart/Item/Delete.php есть активные вызовы dd(), которые могут остановить выполнение в production.",
        "solution": "Удалить все dd() из production-файлов; заменить на Log::debug() там, где нужна диагностика; ограничить Debugbar/Telescope только локальным окружением.",
        "improvement": "Устранение случайных availability-инцидентов; TTFB −5–10 % за счёт удаления profiling overhead.",
    },
    {
        "problem": "Генерация YML/Excel прайса: вся выборка и весь файл в памяти PHP",
        "cause": "Price/Index.php и Price/Base.php строят весь файл одним $query->get(), загружая все 22 000 товаров в память одного воркера синхронно; это блокирует воркер до завершения.",
        "solution": "Использовать chunkById(500) и response()->streamDownload() для потоковой отдачи; тяжёлые экспорты выносить в queue job с хранением статуса.",
        "improvement": "Пиковое потребление памяти снижается на 80–95 %; HTTP-воркер освобождается немедленно; стабильная работа при любом объёме каталога.",
    },
    {
        "problem": "Отсутствие индексов для частых фильтров и JOIN",
        "cause": "В дампе нет составных и одиночных индексов по items(category_id_1c, type), items(id_1c), items(vendor_code), items(barcode), categories(parent_uuid, is_hide, default_sort), sales(uuid_contractor, id). Фильтрующие запросы делают full/range scan.",
        "solution": "Добавить недостающие индексы через миграцию; перед применением в production проверить EXPLAIN на staging-БД.",
        "improvement": "Время фильтрующих запросов каталога −30–60 % по rows examined; меньше нагрузки на DB CPU при пиковом трафике.",
    },
    {
        "problem": "Frontend: отсутствие vite.config.js, глобальное подключение тяжёлых JS-модулей",
        "cause": "В корне проекта нет vite.config.js, поэтому production build нестабилен. Layout подключает jquery-ui, inputmask, fancybox глобально, даже на страницах, где они не нужны.",
        "solution": "Создать vite.config.js с раздельными entrypoints (public, private, catalogue, admin); подключать бандлы через @stack и @push только на нужных страницах.",
        "improvement": "JS/CSS payload −40–60 %; FCP/LCP −20–35 %; TTI улучшается за счёт меньшего объёма парсинга скриптов.",
    },
    {
        "problem": "Отсутствие gzip/brotli сжатия и долгосрочного кэша статических файлов",
        "cause": "В public/.htaccess и nginx-конфигурации нет правил gzip/brotli и Cache-Control: immutable для hashed assets, поэтому каждый визит загружает статику заново.",
        "solution": "Добавить в nginx: gzip/brotli, Cache-Control: public, max-age=31536000, immutable для /build/, Cache-Control для fonts/images; настроить HTTP/2.",
        "improvement": "FCP/LCP −15–25 % на первом визите; для повторных визитов −50–80 % трафика из origin за счёт браузерного кэша.",
    },
    {
        "problem": "Шрифты в устаревших форматах (TTF/EOT/OTF/SVG) без font-display: swap",
        "cause": "В resources/fonts подключены тяжёлые форматы без WOFF2-subset стратегии и без font-display: swap, что блокирует рендер текста до загрузки шрифта.",
        "solution": "Конвертировать в WOFF2 с subset кириллица/латиница, оставить только нужные начертания, добавить font-display: swap и preload для критичного шрифта.",
        "improvement": "FCP −5–15 %; CLS → 0; объём загрузки шрифтов −60–80 %; устранение блокировки рендера текста.",
    },
]

HEADERS = [
    "Проблема",
    "Что приводит к проблеме",
    "Возможное решение",
    "Потенциальное улучшение (показания)",
]

COL_WIDTHS = [Cm(3.8), Cm(4.8), Cm(5.2), Cm(4.4)]


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_para(cell, text: str, bold: bool = False, font_size: int = 10, color=None):
    para = cell.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para


def build_docx(output_path: str):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)

    # Title
    title = doc.add_heading("Анализ проблем производительности Laravel-приложения", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        "Таблица подготовлена в рамках решения тестового задания по оптимизации "
        "веб-приложения на Laravel (анализ проекта B2B-каталога)."
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    # Table
    num_rows = len(ROWS) + 1  # +1 header
    table = doc.add_table(rows=num_rows, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = COL_WIDTHS[i]

    # Header row
    header_row = table.rows[0]
    for i, header_text in enumerate(HEADERS):
        cell = header_row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, "1F4E79")
        # Clear default paragraph
        for p in cell.paragraphs:
            p.clear()
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row_idx, data in enumerate(ROWS):
        row = table.rows[row_idx + 1]
        bg_color = "D6E4F0" if row_idx % 2 == 0 else "FFFFFF"

        for col_idx, key in enumerate(["problem", "cause", "solution", "improvement"]):
            cell = row.cells[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_bg(cell, bg_color)
            for p in cell.paragraphs:
                p.clear()
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(data[key])
            run.font.size = Pt(9)
            if col_idx == 0:
                run.bold = True

    # Footer note
    doc.add_paragraph()
    note = doc.add_paragraph(
        "Источник анализа: статический аудит кодовой базы Laravel (архив 1c8_private-master.zip) "
        "и дампа базы данных MySQL (alfastok_db_1c8_8.sql). "
        "Все метрики улучшений — расчётные оценки на основе объёма данных в таблицах "
        "и характеристик выявленных паттернов; подтверждение требует замеров в production (APM, Lighthouse, slow query log)."
    )
    for run in note.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(120, 120, 120)

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    import os
    out = os.path.join(os.path.dirname(__file__), "..", "speed_opt", "performance_analysis.docx")
    build_docx(os.path.abspath(out))
